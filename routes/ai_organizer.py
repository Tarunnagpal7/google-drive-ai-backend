# ai_organizer.py
import os
import json
from flask import Blueprint, request, jsonify, session
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2.credentials import Credentials
from openai import OpenAI
import re
from typing import List, Dict, Any

import fitz  # PyMuPDF
from docx import Document
import tempfile
def extract_file_content(service, file_id, mime_type) -> str:
    """
    Extracts first ~200 words from PDF or DOCX.
    Skips unsupported types like images/videos.
    Limits extraction to keep processing fast even with 100+ files.
    """
    try:
        metadata = service.files().get(fileId=file_id, fields="size, name").execute()
        if int(metadata.get('size', 0)) > 10_000_000:
            print(f"Skipping large file: {metadata['name']}")
            return ""

        content = ""

        with tempfile.NamedTemporaryFile(delete=False, suffix=".bin") as tmp_file:
            tmp_path = tmp_file.name
            request = service.files().get_media(fileId=file_id)
            downloader = MediaIoBaseDownload(tmp_file, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()

        if mime_type == 'application/pdf':
            pdf_path = tmp_path.replace(".bin", ".pdf")
            os.rename(tmp_path, pdf_path)
            doc = fitz.open(pdf_path)

            words = []
            for page in doc:
                text = page.get_text()
                words.extend(text.split())
                if len(words) >= 200:
                    break
            doc.close()
            os.remove(pdf_path)
            content = " ".join(words[:200])

        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            docx_path = tmp_path.replace(".bin", ".docx")
            os.rename(tmp_path, docx_path)
            doc = Document(docx_path)

            words = []
            for para in doc.paragraphs:
                words.extend(para.text.split())
                if len(words) >= 200:
                    break
            os.remove(docx_path)
            content = " ".join(words[:200])

        else:
            os.remove(tmp_path)
            return ""  # Skip non-text files (image, video, etc.)

        return content.strip()

    except Exception as e:
        print(f"[extract_file_content] Error for {file_id}: {e}")
        return ""
    """
    Extracts first ~200 words from PDF or DOCX.
    Skips unsupported types like images/videos.
    Limits extraction to keep processing fast even with 100+ files.
    """
    try:
        # Download the file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=".bin") as tmp_file:
            tmp_path = tmp_file.name
            request = service.files().get_media(fileId=file_id)
            file_data = request.execute()
            tmp_file.write(file_data)

        content = ""

        if mime_type == 'application/pdf':
            pdf_path = tmp_path.replace(".bin", ".pdf")
            os.rename(tmp_path, pdf_path)
            doc = fitz.open(pdf_path)

            words = []
            for page in doc:
                text = page.get_text()
                words.extend(text.split())
                if len(words) >= 200:
                    break
            doc.close()
            os.remove(pdf_path)
            content = " ".join(words[:200])

        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            docx_path = tmp_path.replace(".bin", ".docx")
            os.rename(tmp_path, docx_path)
            doc = Document(docx_path)

            words = []
            for para in doc.paragraphs:
                words.extend(para.text.split())
                if len(words) >= 200:
                    break
            os.remove(docx_path)
            content = " ".join(words[:200])

        else:
            os.remove(tmp_path)
            return ""  # Skip non-text files (image, video, etc.)

        return content.strip()

    except Exception as e:
        print(f"[extract_file_content] Error for {file_id}: {e}")
        return ""

ai_bp = Blueprint("ai_organizer", __name__)

# Initialize OpenAI client
openai_api = os.environ.get("OPENAI_KEY")
client = OpenAI(api_key=openai_api)


def generate_file_rename_prompt(file_list: List[Dict[str, Any]], pattern: str) -> str:
    """Generate a detailed prompt for AI file organization"""

    files_text = ""
    for i, file in enumerate(file_list, 1):
        files_text += f"{i}. {file['name']} ({file['type']})\n"    

    prompt = f"""
You are a professional AI assistant and expert file organizer. I have {len(file_list)} files and folders that need to be cleaned, renamed, and optionally organized into a more meaningful folder structure.

**Naming Instructions:**
- Use this optional naming format if provided by the user: {pattern or 'No custom pattern given'}
- Follow common naming conventions like PascalCase or snake_case
- Keep file extensions intact
- Remove unnecessary characters, underscores, numbers, or redundant words
- Avoid duplicate or conflicting names
- Ensure clarity, readability, and relevance to content

**Foldering Instructions:**
- Group related files under folders such as Reports, Invoices, Projects, Chapters, etc.
- If a company or client name is detected (like Spacenos, Sensewire, Infiprime, RIT, OmmRudraksha), use it as the top-level folder.
- Do **not** move a folder into itself or into another folder with the same name — skip such suggestions.
- Avoid redundant nesting (e.g., "OmmRudraksha/OmmRudraksha_Reports")
- If no business context is clear, use "General" or skip foldering

**Required Output (JSON only):**
For every file or folder, suggest:
- A cleaner, more professional new name
- A logical folder path (if reorganization is useful)
- A reason explaining why the suggestion was made

Respond strictly in the following JSON format:

[
  {{
    "id": "original_file_id",
    "currentName": "original_name.ext",
    "newName": "Improved_File_Name.ext",
    "newFolder": "Company/Department" (optional),
    "type": "file" or "folder",
    "reason": "Short explanation for the change"
  }}
]

Only return valid JSON. Do not include markdown, commentary, or any additional text.
Here is the list of files/folders:

{files_text}
"""

    return prompt

def get_ai_suggestions(prompt: str) -> List[Dict[str, Any]]:
    """Get AI suggestions for file organization"""
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system", 
                    "content": "You are a professional file organizer. Always respond with valid JSON only."
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        
        content = response.choices[0].message.content.strip()
        
        # Clean up the response to ensure it's valid JSON
        if content.startswith("```json"):
            content = content[7:-3]
        elif content.startswith("```"):
            content = content[3:-3]
        
        suggestions = json.loads(content)
        return suggestions
        
    except json.JSONDecodeError as e:
        print(f"JSON decode error: {e}")
        print(f"Response content: {content}")
        return []
    except Exception as e:
        print(f"OpenAI API error: {e}")
        return []

def create_folder_if_not_exists(service, folder_name: str, parent_id: str = None) -> str:
    """Create a folder in Google Drive if it doesn't exist"""
    try:
        # Check if folder already exists
        query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
        if parent_id:
            query += f" and '{parent_id}' in parents"
        
        results = service.files().list(q=query).execute()
        items = results.get('files', [])
        
        if items:
            return items[0]['id']
        
        # Create new folder
        folder_metadata = {
            'name': folder_name,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        
        if parent_id:
            folder_metadata['parents'] = [parent_id]
        
        folder = service.files().create(body=folder_metadata, fields='id').execute()
        return folder.get('id')
        
    except Exception as e:
        print(f"Error creating folder {folder_name}: {e}")
        return None

def create_nested_folders(service, folder_path: str, parent_id: str = None) -> str:
    """Create nested folders based on path like 'Course_Materials/Chapters'"""
    if not folder_path:
        return parent_id
    
    folders = folder_path.split('/')
    current_parent = parent_id
    
    for folder_name in folders:
        folder_name = folder_name.strip()
        if folder_name:
            current_parent = create_folder_if_not_exists(service, folder_name, current_parent)
            if not current_parent:
                return None
    
    return current_parent

@ai_bp.route('/ai/rename-preview', methods=['POST'])
def ai_rename_preview():
    try:
        if 'credentials' not in session:
            return jsonify({"error": "Not authorized"}), 401

        data = request.get_json()
        pattern = data.get('pattern', "")
        selected_files = data.get('selectedFiles', [])
        if not selected_files:
            return jsonify({"error": "No files selected"}), 400

        creds = Credentials(**session['credentials'])
        service = build('drive', 'v3', credentials=creds)

        enriched_files = []
        for f in selected_files:
            file_id = f['id']
            mime_type = service.files().get(fileId=file_id, fields="mimeType").execute()['mimeType']
            content = extract_file_content(service, file_id, mime_type) if f['type'] == 'file' else ""
            enriched_files.append({
                "id": file_id,
                "name": f['name'],
                "type": f['type'],
                "content": content
            })

        prompt = generate_file_rename_prompt(enriched_files, pattern)
        suggestions = get_ai_suggestions(prompt)

        for i, suggestion in enumerate(suggestions):
            if i < len(enriched_files):
                suggestion['id'] = enriched_files[i]['id']

        return jsonify(suggestions)
    except Exception as e:
        print(f"Error in ai_rename_preview: {e}")
        return jsonify({"error": str(e)}), 500

    """Generate AI suggestions for file renaming and organization"""
    try:
        if 'credentials' not in session:
            return jsonify({"error": "Not authorized"}), 401
        
        data = request.get_json()
        selected_files = data.get('selectedFiles', [])
        
        if not selected_files:
            return jsonify({"error": "No files selected"}), 400
        
        # Generate AI prompt
        prompt = generate_file_rename_prompt(selected_files)
        
        # Get AI suggestions
        suggestions = get_ai_suggestions(prompt)
        
        if not suggestions:
            return jsonify({"error": "Failed to generate AI suggestions"}), 500
        
        # Add original IDs to suggestions
        for i, suggestion in enumerate(suggestions):
            if i < len(selected_files):
                suggestion['id'] = selected_files[i]['id']
        
        return jsonify(suggestions)
        
    except Exception as e:
        print(f"Error in ai_rename_preview: {e}")
        return jsonify({"error": str(e)}), 500

@ai_bp.route('/ai/execute-rename', methods=['POST'])
def execute_rename():
    """Execute the AI suggestions to rename and move files, and delete empty folders"""
    try:
        if 'credentials' not in session:
            return jsonify({"error": "Not authorized"}), 401

        data = request.get_json()
        suggestions = data.get('suggestions', [])

        if not suggestions:
            return jsonify({"error": "No suggestions provided"}), 400

        creds = Credentials(**session['credentials'])
        service = build('drive', 'v3', credentials=creds)

        results = []
        touched_folders = set()

        for suggestion in suggestions:
            try:
                file_id = suggestion['id']
                new_name = suggestion['newName']
                new_folder_path = suggestion.get('newFolder')
                update_metadata = {'name': new_name}

                file_info = service.files().get(fileId=file_id, fields='parents').execute()
                previous_parents = file_info.get('parents', [])
                if previous_parents:
                    touched_folders.update(previous_parents)

                if new_folder_path:
                    folder_id = create_nested_folders(service, new_folder_path)

                    if folder_id:
                        service.files().update(
                            fileId=file_id,
                            body=update_metadata,
                            addParents=folder_id,
                            removeParents=','.join(previous_parents) if previous_parents else None,
                            fields='id, name, parents'
                        ).execute()

                        results.append({
                            'id': file_id,
                            'status': 'success',
                            'message': f'Renamed to \"{new_name}\" and moved to \"{new_folder_path}\"'
                        })
                    else:
                        service.files().update(
                            fileId=file_id,
                            body=update_metadata,
                            fields='id, name'
                        ).execute()

                        results.append({
                            'id': file_id,
                            'status': 'partial',
                            'message': f'Renamed to \"{new_name}\" but folder creation failed'
                        })
                else:
                    service.files().update(
                        fileId=file_id,
                        body=update_metadata,
                        fields='id, name'
                    ).execute()

                    results.append({
                        'id': file_id,
                        'status': 'success',
                        'message': f'Renamed to \"{new_name}\"'
                    })

            except Exception as e:
                results.append({
                    'id': suggestion.get('id', 'unknown'),
                    'status': 'error',
                    'message': f'Error: {str(e)}'
                })

        # 🔥 Delete empty folders
        for folder_id in touched_folders:
            try:
                children = service.files().list(
                    q=f"'{folder_id}' in parents and trashed=false",
                    fields='files(id)',
                    pageSize=1
                ).execute().get('files', [])
                
                if not children:
                    service.files().delete(fileId=folder_id).execute()
                    print(f"✅ Deleted empty folder: {folder_id}")
            except Exception as e:
                print(f"⚠️ Failed to check/delete folder {folder_id}: {e}")

        return jsonify({
            'success': True,
            'results': results,
            'total': len(suggestions),
            'successful': len([r for r in results if r['status'] == 'success']),
            'failed': len([r for r in results if r['status'] == 'error'])
        })

    except Exception as e:
        print(f"Error in execute_rename: {e}")
        return jsonify({"error": str(e)}), 500


@ai_bp.route('/ai/batch-organize', methods=['POST'])
def batch_organize():
    """Organize all files in a folder with AI suggestions"""
    try:
        if 'credentials' not in session:
            return jsonify({"error": "Not authorized"}), 401
        
        data = request.get_json()
        folder_id = data.get('folderId')
        
        if not folder_id:
            return jsonify({"error": "No folder ID provided"}), 400
        
        creds = Credentials(**session['credentials'])
        service = build('drive', 'v3', credentials=creds)
        
        # Get all files in the folder
        results = service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields='files(id, name, mimeType)'
        ).execute()
        
        files = results.get('files', [])
        
        if not files:
            return jsonify({"error": "No files found in folder"}), 400
        
        # Convert to expected format
        selected_files = []
        for file in files:
            file_type = 'folder' if file['mimeType'] == 'application/vnd.google-apps.folder' else 'file'
            selected_files.append({
                'id': file['id'],
                'name': file['name'],
                'type': file_type
            })
        
        # Generate AI suggestions
        prompt = generate_file_rename_prompt(selected_files)
        suggestions = get_ai_suggestions(prompt)
        
        if not suggestions:
            return jsonify({"error": "Failed to generate AI suggestions"}), 500
        
        # Add original IDs to suggestions
        for i, suggestion in enumerate(suggestions):
            if i < len(selected_files):
                suggestion['id'] = selected_files[i]['id']
        
        return jsonify({
            'suggestions': suggestions,
            'total_files': len(files)
        })
        
    except Exception as e:
        print(f"Error in batch_organize: {e}")
        return jsonify({"error": str(e)}), 500

